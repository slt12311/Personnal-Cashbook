#报告下载

from tkinter import *
from PIL import ImageTk, Image
import tkinter.messagebox
from tkinter import filedialog, dialog
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_LINE_SPACING
from docx.oxml.ns import qn
import Function
import matplotlib
matplotlib.use('agg')
import matplotlib.pyplot as plt
import Analyze

#报告下载
class DownLoad:
    #返回控制
    def b_back(self):
        self.root.destroy()
        self.goMain(user=self.user)

    # 生成报告文档
    #预算
    def Budget(self,document):
        # 获取预算
        sql = "select budget from Budget where id='%s'" % self.user
        self.con.execute(sql)
        budgetget = self.con.fetchall()
        a = document.add_paragraph()
        run = a.add_run('预算')
        run.bold = True
        run.font.color.rgb = RGBColor(192, 30, 119)

        #根据预算情况，写入文件
        if (len(budgetget) != 0):
            budget = budgetget[0][0]
            sum = Analyze.Budget.getSum(self=self)
            t = sum[0] / budget
            t = round(t, 2)
            text = ''
            if (sum[0] <= budget):
                text = '''当前月预算共:￥%.2f\n月预算剩余:￥%.2f\n总支出:￥%.2f\n总收入:￥%.2f''' % (budget, budget - sum[0], sum[0], sum[1])
            else:
                text = '''当前月预算共:￥%.2f\n当前月预算超支:￥%.2f\n总支出:￥%.2f\n总收入:￥%.2f''' % (
                budget, sum[0] - budget, sum[0], sum[1])
            b = document.add_paragraph(text)
            b.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            a.add_run('当前未设置预算')
        a.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #统计图
    def Statistics(self,document):
        a = document.add_paragraph()
        run = a.add_run('\n分类统计')
        run.bold = True
        run.font.color.rgb = RGBColor(192, 30, 119)
        a.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        time2 = self.time[0] + '-' + self.time[1] + '-'
        temp = Analyze.Statistics.getSta(self=self, time=time2)
        result2 = []

        #收入
        if (temp == 2 or temp == 4):
            a = document.add_paragraph()
            a.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            a.add_run('收入').bold = True
            result0 = Analyze.Statistics.getIncome(self=self, tag=1, result2=result2, time_b=self.time)
            pie = Analyze.Statistics.Show(self=self, result0=result0, result2=result2, size=3, a=0)
            a = document.add_paragraph()
            a.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = a.add_run('')
            run.add_picture('./%s/cat.jpg' % self.table)

            sum = 0
            for i in result2:
                sum += i
            text = "总收入：￥%.2f\n最高收入来自%s" % (sum, (result0[result2.index(max(result2))]))
            a = document.add_paragraph(text)
            a.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        #支出
        result2 = []
        if (temp == 3 or temp == 4):
            a = document.add_paragraph()
            a.add_run('支出').bold = True
            a.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            result0 = Analyze.Statistics.getIncome(self=self, tag=0, result2=result2, time_b=self.time)
            pie = Analyze.Statistics.Show(self=self, result0=result0, result2=result2, size=3, a=1)

            sum = 0
            for i in result2:
                sum += i
            text = "总支出：￥%.2f\n最高支出来自%s" % (sum, (result0[result2.index(max(result2))]))
            a = document.add_paragraph()
            a.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = a.add_run('')
            run.add_picture('./%s/cat.jpg' % self.table)
    #时间折线图
    def TimeSum(self,document):
        a = document.add_paragraph()
        run = a.add_run('\n收入/支出时间趋势图')
        run.bold = True
        run.font.color.rgb = RGBColor(192, 30, 119)
        a.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # 统计
        self.table_list = []
        Function.Check.allRecord(self=self)
        j = 0
        sum1 = []
        sum2 = []
        name=[]
        #获取五期或以下月收支数据
        for i in self.table_list:
            if (j < 5):
                sql = "select sum(money) from [%s] where time like '%%%s-%%' and prime='收入'" % (self.table, i)
                self.con.execute(sql)
                result = self.con.fetchall()
                if (result[0][0] == None):
                    result[0][0] = 0
                sum1.append(result[0][0])

                sql = "select sum(money) from [%s] where time like '%%%s-%%' and prime='支出'" % (self.table, i)
                self.con.execute(sql)
                result = self.con.fetchall()
                if (result[0][0] == None):
                    result[0][0] = 0
                sum2.append(result[0][0])
                name.append(self.table_list[j])
            else:
                break
            j += 1
        name_b=[]
        sum1_b=[]
        sum2_b=[]
        len1=len(name)
        for i in range(len1):
            name_b.append(name[len1-i-1])
            sum1_b.append(sum1[len1-i-1])
            sum2_b.append(sum2[len1-i-1])
        # 绘图
        fig = plt.figure(figsize=(6, 4))
        plt.plot(name_b,  # x轴数据
                 sum1_b,  # y轴数据
                 linestyle='-',  # 折线类型
                 linewidth=2,  # 折线宽度
                 color='steelblue',  # 折线颜色S
                 marker='o',  # 点的形状
                 markersize=6,  # 点的大小
                 markeredgecolor='black',  # 点的边框色
                 markerfacecolor='steelblue',  # 点的填充色
                 label='收入')  # 添加标签
        plt.plot(name_b,  # x轴数据
                 sum2_b,  # y轴数据
                 linestyle='-',  # 折线类型
                 linewidth=2,  # 折线宽度
                 color='#ff9999',  # 折线颜色
                 marker='o',  # 点的形状
                 markersize=6,  # 点的大小
                 markeredgecolor='black',  # 点的边框色
                 markerfacecolor='#ff9999',  # 点的填充色
                 label='支出')  # 添加标签
        # print(name,sum1,sum2)
        plt.tick_params(top='off', right='off')
        if(j==1):
            plt.title('%s收入/支出统计图' % self.table_list[0])
        else:
            plt.title('%s至%s收入/支出统计图' % (self.table_list[j - 1], self.table_list[0]))
        plt.xlabel('日期')
        plt.ylabel('金额')
        plt.legend()
        plt.savefig('./%s/Cat.jpg' % self.table)
        plt.close()
        #写入文档
        a = document.add_paragraph()
        a.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = a.add_run('')
        run.add_picture('./%s/cat.jpg' % self.table)
    #制作文档
    def Make(self):

        # 文件
        document = Document()
        #标题
        time0 = self.time[0] + '年' + self.time[1] + '月'
        head = document.add_heading(0)
        name=document.add_paragraph(0)
        run = head.add_run('%s 财务报告'% time0)
        head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run=name.add_run('用户:%s'%self.user)
        name.style='Intense Quote'
        name.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        #预算
        self.Budget(document=document)

        #分类模块
        self.Statistics(document=document)

        #时间统计模块
        self.TimeSum(document=document)
        
        #调整保存
        document.styles['Normal'].font.name = 'Times New Roman'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        try:
            document.save('%s/%s.docx'%(self.table,self.user))
        except IOError:
            tkinter.messagebox.showwarning("警告","请关闭您已经打开的%s.docx文件"%self.user)
    #文件保存
    def Down(self):
        #制作文档
        self.Make()
        #下载
        file_path = filedialog.asksaveasfilename(title=u'保存文件',filetypes=[('Word文档', '*.docx')])
        if (file_path != ''):
            name=file_path.split('/')[-1]
            if(name[-5:]=='.docx'):
                file_path=file_path[:-5]
                name=name[:-5]
            if (name.isalnum()!=TRUE):
                tkinter.messagebox.showwarning("警告", "文件名不可以含有特殊字符")
            else:
                document=Document('%s/%s.docx'%(self.table,self.user))
                file_path+='.docx'
                try:
                    document.save(file_path)
                    tkinter.messagebox.showinfo("成功","保存成功！")
                except IOError:
                    tkinter.messagebox.showwarning("警告", "请关闭您已经打开的%s.docx文件" % file_path)

    #初始化/界面
    def __init__(self, con, user, goMain, time1):
        self.con = con
        self.user = user
        self.table = 'user' + user
        self.goMain = goMain
        self.time = time1
        time0 = time1[0] + '年' + time1[1] + '月'

        # 界面初始
        self.root = Tk()
        Analyze.goHelp(root=self.root)
        self.root.title("报告下载")
        self.root.resizable(0, 0)
        self.root.geometry('500x300')
        image1 = Image.open('source/鸟底纹素材.png')
        image2 = Image.open(r'source/鸟底纹素材2.png')
        image3 = Image.open(r'source/鸟底纹素材2.jpg')
        back_image2 = ImageTk.PhotoImage(image3)
        background_image = ImageTk.PhotoImage(image2)
        back_image = ImageTk.PhotoImage(image1)
        Label(self.root, image=back_image).place(x=0, y=0)
        label = Label(self.root, text='%s 报告下载' % time0, compound='center',
                      image=background_image)
        label.place(relx=0.15, rely=0.05)
        Button(self.root, text="报告下载", compound='center', command=self.Down, image=back_image2).place(relx=0.34,
                                                                                                      rely=0.3)
        Button(self.root, text='返回', command=self.b_back, bg='snow').place(relx=0.45, rely=0.7)
        self.root.mainloop()

